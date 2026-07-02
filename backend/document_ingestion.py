import contextlib
import io
import os
import re

import fitz  # PyMuPDF
import pytesseract
from PIL import Image

from audio_processor import AudioProcessorService
from file_validator import FileValidator
from memory_extractor_service import MemoryExtractorService


class DocumentExtractor:
    """Extract and normalize text from various document formats with robust error handling."""

    @staticmethod
    def normalize_text(text: str) -> str:
        """Normalize extracted text: clean up whitespace, line endings, and control characters."""
        if not text:
            return ""
        # Replace carriage returns
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        # Remove non-printable control characters but keep standard whitespace/newlines
        text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uD7FF\uE000-\uFFFF]", "", text)
        # Replace multiple spaces/tabs with a single space
        text = re.sub(r"[ \t]+", " ", text)
        # Normalize paragraph spacing (maximum 2 consecutive newlines)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF file using PyMuPDF, falling back to OCR for scanned pages."""
        text = ""
        doc = None
        try:
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc):
                page_text = page.get_text()

                # If page contains little to no digital text, attempt OCR
                if len(page_text.strip()) < 50:
                    try:
                        # Render page to image at 150 DPI for OCR
                        pix = page.get_pixmap(dpi=150)
                        img_data = pix.tobytes("png")
                        img = Image.open(io.BytesIO(img_data))
                        ocr_text = pytesseract.image_to_string(img)
                        if ocr_text.strip():
                            page_text = ocr_text
                    except Exception as ocr_err:
                        print(f"OCR failed on PDF page {page_num}: {ocr_err}")

                text += page_text + "\n"

            return DocumentExtractor.normalize_text(text)
        except Exception as e:
            raise RuntimeError(f"Failed to extract text from PDF: {e!s}")
        finally:
            if doc:
                with contextlib.suppress(BaseException):
                    doc.close()

    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """Extract text from TXT file with encoding detection fallback."""
        encodings = ["utf-8", "latin-1", "cp1252", "utf-16"]
        for encoding in encodings:
            try:
                with open(file_path, encoding=encoding) as f:
                    text = f.read()
                return DocumentExtractor.normalize_text(text)
            except UnicodeDecodeError:
                continue
            except Exception as e:
                raise RuntimeError(f"Failed to read TXT file: {e!s}")

        raise RuntimeError("Failed to decode TXT file with any supported encoding.")

    @staticmethod
    def extract_text_from_image(file_path: str) -> str:
        """Extract text from image file using Tesseract OCR."""
        try:
            # Prevent decompression bomb errors on large images
            Image.MAX_IMAGE_PIXELS = None
            with Image.open(file_path) as image:
                text = pytesseract.image_to_string(image)
            return DocumentExtractor.normalize_text(text)
        except Exception as e:
            raise RuntimeError(f"Failed to extract text from image: {e!s}")

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX/DOC file using python-docx."""
        try:
            from docx import Document

            doc = Document(file_path)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            # Also extract table cell text
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)
            return DocumentExtractor.normalize_text("\n".join(parts))
        except ImportError:
            raise RuntimeError("python-docx is not installed. Run: pip install python-docx")
        except Exception as e:
            raise RuntimeError(f"Failed to extract text from DOCX: {e!s}")

    @staticmethod
    def extract_text_from_plain(file_path: str) -> str:
        """Extract text from plain-text formats: md, csv, json, xml, html, rtf, log, code files."""
        encodings = ["utf-8", "latin-1", "cp1252", "utf-16"]
        for encoding in encodings:
            try:
                with open(file_path, encoding=encoding, errors="replace") as f:
                    return DocumentExtractor.normalize_text(f.read())
            except Exception:
                continue
        raise RuntimeError("Failed to decode file with any supported encoding.")

    @staticmethod
    def extract_text(
        file_path: str,
        file_type: str,
        audio_processor: AudioProcessorService | None = None,
    ) -> str:
        """
        Extract text from a file based on its type with graceful error handling.

        Args:
            file_path: Path to the file
            file_type: File type (pdf, txt, png, jpg, jpeg, wav, mp3, etc.)
            audio_processor: AudioProcessorService for audio transcription

        Returns:
            Extracted and normalized text
        """
        file_typelower = file_type.lower()

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        # Validate file signature for security (except txt files)
        if file_typelower != "txt":
            if not FileValidator.validate_file_type(file_path, file_typelower):
                detected_type = FileValidator.get_detected_type(file_path)
                error_msg = f"File signature does not match extension. Expected: {file_typelower}"
                if detected_type:
                    error_msg += f", Detected: {detected_type}"
                raise ValueError(error_msg)

        # Limit file size processing on CPU to avoid hanging (e.g., max 50MB for docs, 100MB for audio)
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        if file_typelower in ["pdf", "txt", "png", "jpg", "jpeg"] and file_size_mb > 50:
            raise ValueError(
                f"File size ({file_size_mb:.1f}MB) exceeds the 50MB CPU processing limit."
            )

        try:
            if file_typelower == "pdf":
                return DocumentExtractor.extract_text_from_pdf(file_path)
            elif file_typelower in [
                "txt",
                "md",
                "csv",
                "json",
                "xml",
                "html",
                "log",
                "py",
                "js",
                "ts",
                "jsx",
                "tsx",
                "java",
                "c",
                "cpp",
                "cs",
                "go",
                "rs",
                "rtf",
                "odt",
            ]:
                return DocumentExtractor.extract_text_from_plain(file_path)
            elif file_typelower in ["docx", "doc"]:
                return DocumentExtractor.extract_text_from_docx(file_path)
            elif file_typelower in ["png", "jpg", "jpeg", "gif", "bmp", "tiff"]:
                return DocumentExtractor.extract_text_from_image(file_path)
            elif file_typelower in ["wav", "mp3"]:
                if audio_processor is None:
                    raise RuntimeError("Audio processor required for audio transcription")
                return audio_processor.transcribe_audio(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            print(f"Extraction error for {file_path} ({file_type}): {e}")
            raise e


class MemoryExtractor:
    """Extract memories from document text using local LLM."""

    def __init__(self, llm):
        """
        Initialize memory extractor with local LLM.

        Args:
            llm: LocalLLM instance
        """
        self.llm = llm
        self.structured_extractor = MemoryExtractorService(llm)

    def extract_memories(
        self, text: str, source_document: str | None = None, max_memories: int = 5
    ) -> list:
        """
        Extract key memories from document text.

        Args:
            text: Document text
            source_document: Name of the source document
            max_memories: Maximum number of memories to extract

        Returns:
            List of extracted memories (title, content, tags, source_document)
        """
        if not text or len(text.strip()) < 20:
            return []

        try:
            # Use the consolidated structured extractor
            structured_memories = self.structured_extractor.extract_memories(
                text, source_document, max_memories
            )

            # Convert to list format for backward compatibility
            memories = []
            for memory_schema in structured_memories:
                memories.append(
                    {
                        "title": memory_schema.title,
                        "content": memory_schema.summary,
                        "tags": memory_schema.type,
                        "source_document": source_document,
                        "structured_data": memory_schema,  # Include full structured data
                    }
                )

            return memories

        except Exception as e:
            print(f"Error extracting memories: {e}")
            return []

    def extract_structured_memories(
        self, text: str, source_document: str | None = None, max_memories: int = 5
    ) -> list:
        """
        Extract structured memories from document text.

        Args:
            text: Document text
            source_document: Name of the source document
            max_memories: Maximum number of memories to extract

        Returns:
            List of MemorySchema objects
        """
        return self.structured_extractor.extract_memories(text, source_document, max_memories)
